from docx import Document
from pathlib import Path

SAMPLE_FILE_PATH = Path("lfs/documents/paragraphs.txt")
WRITE_FILE_PATH = Path("lfs/documents/demo.docx")


def create_document() -> Document:
    document = Document()
    document.add_heading('Sample Texts', 0)
    return document


def read_paragraphs(path: Path) -> list:
    with open(path) as f:
        lines = f.readlines()
        return [item for item in lines]


def add_paragraphs_to_doc(doc: Document, par: list) -> Document:
    for item in par:
        doc.add_paragraph(item)
    return doc


def save_document(doc: Document, path: Path):
    doc.save(path)


def main():
    doc = create_document()
    pars = read_paragraphs(SAMPLE_FILE_PATH)
    add_paragraphs_to_doc(doc, pars)
    save_document(doc, WRITE_FILE_PATH)


if __name__ == "__main__":
    main()
