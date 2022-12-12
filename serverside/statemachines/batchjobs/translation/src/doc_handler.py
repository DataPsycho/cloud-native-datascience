from pathlib import Path

import docx
from src.logger import get_logger

LOGGER = get_logger()


def read_document_to_paragraphs(path: Path) -> dict:
    """Read a document from the path
    :param path: path to file
    """
    doc = docx.Document(path)
    doc_item = {"title": None, "paragraphs": []}
    for item in doc.paragraphs:
        if item.style.name == 'Title':
            doc_item["title"] = item.text
        else:
            doc_item["paragraphs"].append(item.text)
    LOGGER.info(f"File read successfully from {str(path)}")
    return doc_item


def create_document(doc_item: dict, path: Path):
    """
    Create translated document from translated items
    :param doc_item: Dictionary of translated text
    :param path: Path to write the file
    :return: No Return, a IO operation
    """
    doc = docx.Document()
    if doc_item["title"]:
        heading = doc.add_heading(doc_item["title"], 0)
    else:
        heading = doc.add_heading("---", 0)
    heading.alignment = 1
    for paragraph in doc_item["paragraphs"]:
        doc.add_paragraph(paragraph)
    LOGGER.info(f"File created successfully to {str(path)}")
    doc.save(path)


if __name__ == "__main__":
    file_path = Path("lfs").joinpath("documents", "demo.docx")
    read_document_to_paragraphs(file_path)
